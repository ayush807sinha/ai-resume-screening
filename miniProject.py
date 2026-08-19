import os
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel
import time


load_dotenv()


my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("Api key error")



client = Groq(api_key=my_api_key)

model = "openai/gpt-oss-120b"

jobdescription = """
JOB DESCRIPTION

Software Development Engineer – Full Stack Java

Company: FinTech Technology Company
Location: Bangalore / Remote
Experience: 0–2 Years

About the Role

We are looking for a motivated Full Stack Java Developer to join our engineering team. The candidate will be responsible for developing scalable backend services and responsive web applications while working closely with product and engineering teams.

Responsibilities

• Design, develop, and maintain scalable backend services using Java and Spring Boot.
• Develop and integrate RESTful APIs for web and application services.
• Build responsive and reusable frontend components using React.js.
• Work with relational databases such as MySQL and write efficient SQL queries.
• Develop microservices and implement service-to-service communication.
• Work with message brokers such as RabbitMQ or Kafka for asynchronous communication.
• Write clean, maintainable, and testable code.
• Debug production issues and optimize application performance.
• Participate in code reviews and follow software engineering best practices.
• Collaborate with frontend developers, product managers, and other engineers.

Required Skills

• Strong knowledge of Java and Object-Oriented Programming.
• Good understanding of Spring Boot and Spring Framework.
• Experience developing REST APIs.
• Knowledge of Spring Data JPA and Hibernate.
• Strong SQL and MySQL knowledge.
• Basic understanding of microservices architecture.
• Experience with React.js, JavaScript, HTML, and CSS.
• Understanding of Git and version control.
• Good problem-solving and Data Structures & Algorithms knowledge.

Preferred Skills

• Experience with Spring Cloud.
• Knowledge of Eureka Service Discovery and API Gateway.
• Experience with RabbitMQ or Apache Kafka.
• Familiarity with MongoDB or other NoSQL databases.
• Knowledge of Docker and containerization.
• Understanding of CI/CD pipelines.
• Familiarity with cloud platforms such as AWS or Azure.

Education

• Bachelor's degree in Computer Science, Information Technology, or a related field.
• Freshers and candidates with 0–2 years of relevant experience are encouraged to apply.

"""

class JobD(BaseModel):
    role:str
    required_skills: list[str]
    prefered_skills: list[str]
    minimum_exp:float | None
    educational_requirements: list[str]
    resposibilities: list[str]


JobDSchema = JobD.model_json_schema()

system_prompt = f"""
    You are an expert HR assistant.

    Your job is to analyze job description and extract the structured information from it.
    
    Return ONLY valid JSON matching this schema: {JobDSchema}
    
    IMPORTANT :
    Do NOT return the schema itself.
    Do NOT return fields like "properties" , "tite" or "type".
    Fill the schema with actual extracted information from the job description.

    If minimum experience is not mentioned, return null.
    If information for a list is missing, return an empty list.
    Do not invent information.    
"""


user_prompt = f"""
    Analyze the following Job description: {jobdescription} 
"""


message_System = {
    "role" : "system",
    "content":system_prompt
}


message_user = {
    "role":"user",
    "content":user_prompt
}

responseFormat = {
    "type":"json_object"
}

messages_list = [message_System,message_user]


response = client.chat.completions.create(model=model,messages=messages_list,response_format=responseFormat)

answer = response.choices[0].message.content


raw_json = answer

import json 

job_data =  json.loads(raw_json)

job = JobD(**job_data)


class MatchResult(BaseModel):
    score:float
    details:dict

class Experience(BaseModel):
    company :str | None=None
    role : str | None=None
    duaration : str | None = None
    discrption : str | None = None
    skills_used : list[str] = []


class Resume(BaseModel):
    name :str | None=None
    email : str | None=None
    phone : str | None=None

    total_exp : float | None=None 

    skills : list[str]
    Experiences : list[Experience]
    education : list[str]
    projects : list[str]
    certifications: list[str]


resume_schema = Resume.model_json_schema()


def final_score(job,resume):

    match_schema = MatchResult.model_json_schema()

    prompt = f"""
        You are an HR recruiter. Compare candidates's resume with the job description.

        JOB DESCRIPTION : {job.model_dump_json(indent=2)}

        CADIDATE RESUME : {resume.model_dump_json(indent=2)}

        Return JSON matching this schema : {match_schema}

        Give me :
        - Candidate Name
        - Mathcing Skills
        - Missing Important Skills
        - Wheather Experience Required is met
        - Overall match percentage from 0 to 100
        - A short final verdict 

        Keep the response concise and easy to read.

    """
    mes_promt = {
        "role":"system",
        "content":prompt
    }

    msglist = [mes_promt]

    res_format = {
        "type":"json_object"
    }

    res = client.chat.completions.create(model=model,messages=msglist,response_format=res_format)

    rawanswer = res.choices[0].message.content

    data = json.loads(rawanswer)
    return MatchResult(**data)






def parsedResuem(resume_text):
    sys_prompt= f"""
        You are an expert resume parser.

        Extract information from the resume based on its meaning, not only based on exact section headings.

        Diffrent Resumes cah have diffrent headings.
        
        For example:
        - Experience
        - Professional Experience
        - Work History
        - Employement
        - Internships    
        
        These may all contain relevent experience.

        Skills may also appear in the skills sectiom , work experience, internship or projects.

        Return ONLY valid JSON matching this sehema: {resume_schema}

        IMPORTNAT:
        - Do not invent information
        - If a value is not available, return null  
        - If a list has no information, return an empty list.
        - Include internships inside experience.
        - Extract skills mentioned accross the entire resume.
    """


    user1_prompt = f"""
        Parse the following resume : {resume_text}
    """

    message_sys = {
        "role":"system",
        "content":sys_prompt
    }

    message_user1 = {
        "role":"user",
        "content":user1_prompt
    }

    messageList = [message_sys,message_user1]

    res_format = {
        "type":"json_object"
    }

    res_Client = client.chat.completions.create(model=model,messages=messageList,response_format=res_format)

    resume_res = res_Client.choices[0].message.content

    data = json.loads(resume_res)
    r = Resume(**data)
    return r




from pypdf import PdfReader
from docx import Document

def readPdf(file_path):
    reader = PdfReader(file_path)
    text=""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text = text + page_text + "\n"
    return text


def readDocx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text = text + paragraph.text + "\n"


    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text = text + cell.text +"\n" 

    return text



def readResume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return readPdf(file_path)
    elif file_path.suffix.lower()==".docx":
        return readDocx(file_path)
    else:
        return None



resume_folder = Path("resume")

all_result = []


for file_path in resume_folder.iterdir():

    if file_path.suffix.lower() not in [".pdf",".docx"]:
        continue

    print("\nProcessing: " ,file_path.name)
    resume_text = readResume(file_path)
    parsed_resume = parsedResuem(resume_text)
    time.sleep(5)

    result = final_score(job,parsed_resume)

    time.sleep(5)

    print("Score : " ,result.score)

    time.sleep(5)

    all_result.append({
        "name" : parsed_resume.name,
        "score": result.score,
        "details":result.details
    })

    all_result.sort(
        key = lambda cadidate: cadidate["score"],
        reverse=True
    )

    print(result.details)

